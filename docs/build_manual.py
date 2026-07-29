from __future__ import annotations

from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "交付"
ASSET_DIR = OUT_DIR / "manual-assets"
OUT_DOCX = OUT_DIR / "安全带耐久试验系统_软件使用调试与落地指导手册.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT = "F4F6F9"
GREEN = "14805E"
ORANGE = "B45309"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111827"


def set_font(run, size=11, bold=False, color=BLACK, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    return run


def shade(element, fill):
    shd = element.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        element.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def setup_numbering(doc):
    numbering = doc.part.numbering_part.element

    def add(kind, abstract_id, num_id):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        if kind == "bullet":
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Symbol")
            fonts.set(qn("w:hAnsi"), "Symbol")
            rpr.append(fonts)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(abstract_id))
        num.append(aid)
        numbering.append(num)

    add("bullet", 90, 90)
    add("decimal", 91, 91)


def set_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


_next_num_id = 100


def new_numbering_instance(doc, abstract_id=91):
    global _next_num_id
    num_id = _next_num_id
    _next_num_id += 1
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId")
    aid.set(qn("w:val"), str(abstract_id))
    num.append(aid)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True, color=NAVY)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        set_num(p, 90)
        set_font(p.add_run(item))


def add_compact_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        set_num(p, 90)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        set_font(p.add_run(item), size=9.5)


def add_steps(doc, items):
    num_id = new_numbering_instance(doc)
    for item in items:
        p = doc.add_paragraph(style="Normal")
        set_num(p, num_id)
        set_font(p.add_run(item))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
             color=BLUE if level < 3 else DARK_BLUE)
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, title, text, kind="info"):
    fill, color = {
        "info": ("EAF3FB", DARK_BLUE),
        "warn": ("FFF4E5", ORANGE),
        "danger": ("FDECEC", RED),
        "ok": ("EAF7F1", GREEN),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title + "  "), bold=True, color=color)
    set_font(p.add_run(text), color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, text in enumerate(headers):
        shade(hdr.cells[i]._tc.get_or_add_tcPr(), LIGHT_BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(str(text)), size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_picture(doc, path, caption, alt, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    set_font(cap.add_run(caption), size=9, color=MUTED)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.05
    shade(p._p.get_or_add_pPr(), "F3F4F6")
    set_font(p.add_run(text), size=8.5, name="Consolas", east_asia="Microsoft YaHei")


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(paragraph.add_run("第 "), size=8.5, color=MUTED)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    set_font(paragraph.add_run(" 页"), size=8.5, color=MUTED)


def make_architecture_png(path):
    canvas = Image.new("RGB", (1500, 720), "white")
    draw = ImageDraw.Draw(canvas)
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    bold_path = Path("C:/Windows/Fonts/msyhbd.ttc")
    f_title = ImageFont.truetype(str(bold_path), 38)
    f_head = ImageFont.truetype(str(bold_path), 25)
    f_body = ImageFont.truetype(str(font_path), 21)
    draw.text((55, 32), "安全带耐久试验系统 - 软件与硬件边界", font=f_title, fill=(11, 37, 69))
    boxes = [
        (60, 125, 350, 275, "WinForms 上位机", "界面 / 配方 / 曲线 / 报警\n历史 / 日志 / 诊断"),
        (430, 125, 720, 275, "ITestEngine", "DemoTestEngine\nHardwareTestEngine"),
        (800, 125, 1090, 275, "IHardwarePlatform", "现场独立适配器 DLL\n厂家 SDK 封装"),
    ]
    colors = [(234, 243, 251), (232, 238, 245), (234, 247, 241)]
    for (x1, y1, x2, y2, head, body), fill in zip(boxes, colors):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=(46, 116, 181), width=3)
        draw.text((x1 + 20, y1 + 18), head, font=f_head, fill=(11, 37, 69))
        draw.multiline_text((x1 + 20, y1 + 65), body, font=f_body, fill=(51, 65, 85), spacing=10)
    for x in (365, 735):
        draw.line((x, 200, x + 50, 200), fill=(46, 116, 181), width=5)
        draw.polygon([(x + 50, 200), (x + 35, 190), (x + 35, 210)], fill=(46, 116, 181))
    hardware = [
        (80, 420, 330, 590, "PCIe CAN 卡", "驱动器 / 电机\n命令与状态"),
        (390, 420, 640, 590, "模拟量采集卡", "拉力 / 电流 / 电压\n原始值与时间戳"),
        (700, 420, 950, 590, "安全输入", "急停 / 安全门\n正反限位"),
        (1010, 420, 1420, 590, "SQLite", "Demo 与正式库隔离\n参数 / 配方 / 记录 / 采样点 / 日志"),
    ]
    for x1, y1, x2, y2, head, body in hardware:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=(244, 246, 249), outline=(100, 116, 139), width=2)
        draw.text((x1 + 18, y1 + 18), head, font=f_head, fill=(31, 77, 120))
        draw.multiline_text((x1 + 18, y1 + 68), body, font=f_body, fill=(51, 65, 85), spacing=8)
    draw.line((945, 275, 945, 370), fill=(20, 128, 94), width=5)
    draw.line((205, 370, 1215, 370), fill=(20, 128, 94), width=5)
    for x in (205, 515, 825, 1215):
        draw.line((x, 370, x, 410), fill=(20, 128, 94), width=5)
        draw.polygon([(x, 420), (x - 10, 402), (x + 10, 402)], fill=(20, 128, 94))
    draw.text((55, 650), "原则：UI 不引用厂家 DLL；任一关键设备或安全联锁异常，Production 模式必须禁止启动。", font=f_head, fill=(155, 28, 28))
    canvas.save(path)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(hp.add_run("安全带耐久试验系统"), size=8.5, bold=True, color=NAVY)
    set_font(hp.add_run("    软件使用调试与落地指导手册"), size=8.5, color=MUTED)
    footer = section.footer
    add_page_field(footer.paragraphs[0])
    setup_numbering(doc)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "system-architecture.png"
    make_architecture_png(architecture)

    doc = Document()
    configure_document(doc)

    # Editorial-cover pattern with compact-reference-guide preset.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("工程实施与调试手册"), size=11, bold=True, color=ORANGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("安全带耐久试验系统"), size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("软件使用、硬件接入、现场调试与正式交付指导手册"), size=14, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    set_font(p.add_run("V1.0 · WinForms / .NET 8 / SQLite · Demo + Production 双模式"), size=10.5, bold=True, color=MUTED)
    add_picture(doc, ROOT / "artifacts" / "captures" / "control.png",
                "当前软件 Demo 实际运行界面（非概念图）",
                "安全带耐久试验系统试验控制页面截图", width=6.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    set_font(p.add_run("沈阳艾德瑞自动化有限公司"), size=12, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(f"编制日期：{date.today():%Y-%m-%d}    文档状态：现场实施基线"), size=9.5, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "文档控制与使用边界", 1)
    add_table(doc, ["项目", "内容"], [
        ("适用对象", "电气工程师、C# 上位机开发人员、设备调试人员、设备维护人员"),
        ("软件基线", "DurabilityTestingSystem V1.0；.NET 8 WinForms；SQLite"),
        ("当前能力", "Demo 演示可运行；Production 通用流程、诊断和安全启动门槛已实现"),
        ("待现场补齐", "CAN 卡/采集卡厂家 SDK、驱动器 CAN 协议、传感器量程与标定、安全 I/O 定义"),
        ("正式数据库", "%LocalAppData%\\SeatbeltDurabilitySystem\\durability.db"),
        ("演示数据库", "%LocalAppData%\\SeatbeltDurabilitySystem\\durability-demo.db"),
    ], [2700, 6660], 9)
    add_callout(doc, "结论", "软件已经具备从 Demo 过渡到正式系统的架构，但硬件型号与协议尚未确定，因此任何人都不能声称当前版本已经完成真机控制。正式模式在适配器未配置或自检失败时会阻止试验启动。", "warn")
    add_heading(doc, "目录", 2)
    for line in [
        "1 系统定位与已实现功能  ·  2 总体架构与工作流程  ·  3 正式化前必须确定的硬件资料",
        "4 硬件选型与电气设计建议  ·  5 软件安装、配置与运行模式  ·  6 现场硬件适配器开发",
        "7 CAN 通讯调试  ·  8 模拟量采集、换算与标定  ·  9 安全联锁与保护",
        "10 配方、试验与数据操作  ·  11 分阶段现场调试流程  ·  12 验收测试",
        "13 数据库、备份与恢复  ·  14 故障排查  ·  15 运维、变更与交付清单",
    ]:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        set_font(p.add_run(line), size=9.5)
    doc.add_page_break()

    add_heading(doc, "1 系统定位与已实现功能", 1)
    add_body(doc, "本系统用于控制安全带自带电机执行往复耐久动作，同时采集拉力、电流、电压和可选位置数据。上位机负责配方、状态机、实时曲线、结果保存、历史查询、日志和诊断；驱动功率、安全停机和最终联锁必须由硬件实现。")
    add_heading(doc, "1.1 当前软件完成度", 2)
    add_table(doc, ["模块", "当前状态", "正式现场还需完成"], [
        ("试验控制", "已完成 UI、循环状态机、启停/暂停/复位、曲线", "将通用动作映射到真实驱动器命令"),
        ("参数设置", "已持久化试验、CAN、采集和保护参数", "按最终设备扩展地址、通道和标定系数"),
        ("试验方案", "方案主参数与步骤明细已保存", "确认每个步骤的完成条件和异常分支"),
        ("数据", "汇总记录、采样点、日志、筛选和 CSV 已实现", "确认采样保存频率、报告模板和保留周期"),
        ("诊断", "自检、设备状态、数据库完整性和备份已实现", "由真实适配器提供设备诊断细节"),
        ("硬件", "已定义适配器接口和安全失败模板", "获得 SDK/协议后编写真正的适配器 DLL"),
    ], [1700, 3100, 4560], 8.6)
    add_heading(doc, "1.2 Demo 与 Production 的区别", 2)
    add_table(doc, ["项目", "Demo", "Production"], [
        ("数据来源", "内置模拟器", "真实采集卡与传感器"),
        ("电机命令", "不访问硬件", "CAN 适配器发送并校验应答"),
        ("数据库", "durability-demo.db", "durability.db"),
        ("启动条件", "模拟设备在线", "全部关键设备及安全联锁自检通过"),
        ("用途", "演示、UI 和工艺确认", "现场生产试验"),
    ], [1800, 3300, 4260], 9)
    add_callout(doc, "安全规则", "现场调试时，看到侧栏“PRODUCTION 正式模式”仍不代表已经可运行；必须在“设备诊断”页确认 CAN、模拟量、电机和安全联锁全部在线，且系统摘要允许启动。", "danger")

    add_heading(doc, "2 总体架构与工作流程", 1)
    add_picture(doc, architecture, "图 1  软件与现场硬件的分层边界", "系统架构图")
    add_heading(doc, "2.1 正常试验流程", 2)
    add_steps(doc, [
        "操作员选择试验方案，填写试件编号，核对目标拉力、循环次数、速度、量程和保护值。",
        "系统连接 CAN 卡、采集卡和安全输入，读取驱动器状态并完成启动前自检。",
        "电机执行正向拉伸；到达设定时间、位置或拉力条件后进入负载保持。",
        "保持结束后执行反向回程，原点或反向限位确认后循环计数加一。",
        "采样线程持续读取并保存拉力、电流、电压、位置、循环号和阶段。",
        "到达循环目标后正常停机并保存汇总；通讯、安全或超限异常时立即进入报警状态并请求停机。"
    ])
    add_heading(doc, "2.2 关键源代码入口", 2)
    add_table(doc, ["文件", "作用", "现场修改原则"], [
        ("Program.cs", "装载运行配置、选择 Demo/Production、创建数据库与引擎", "通常不改；只扩展启动参数"),
        ("Infrastructure/HardwareTestEngine.cs", "正式循环时序、采样、保护和状态流转", "工艺时序变化时修改并回归测试"),
        ("Infrastructure/IHardwarePlatform.cs", "厂家硬件统一接口", "保持稳定，避免 UI 依赖厂家 DLL"),
        ("HardwareAdapter.Template", "现场适配器安全失败模板", "复制为 Site 项目后实现全部 TODO"),
        ("Data/AppDatabase.cs", "SQLite 表、读写、备份与完整性检查", "新增字段要做迁移和备份"),
        ("system-profile.json", "模式及适配器 DLL 配置", "现场发布目录单独维护"),
    ], [3000, 3000, 3360], 8.3)

    add_heading(doc, "3 正式化前必须确定的硬件资料", 1)
    add_body(doc, "在硬件未定型前不要写死厂家 API 或 CAN 报文。采购下单前，至少取得以下资料并形成《接口冻结表》。")
    add_table(doc, ["设备", "必须取得的资料", "没有资料的后果"], [
        ("工控机", "CPU/内存/SSD、Windows 版本、PCIe 插槽数量和规格、网口、USB、供电、散热", "CAN/采集卡装不下或驱动不兼容"),
        ("CAN 卡", "准确型号、x64 驱动、.NET/C SDK、示例、通道数、隔离、支持波特率", "C# 无法调用或现场不稳定"),
        ("安全带电机/驱动器", "品牌型号、CAN 协议、节点配置、报文 ID、字节序、比例、心跳、报警表", "无法安全控制电机"),
        ("模拟量采集卡", "接口方式、输入类型、分辨率、采样率、是否同步采样、隔离、SDK/寄存器表", "读不到数据或量程不匹配"),
        ("拉力传感器", "量程、精度、过载、mV/V 或变送输出、供电、标定证书", "无法正确换算 N"),
        ("电流/电压传感器", "被测类型、测量点、量程、输出、隔离、响应时间和方向", "读数无意义或损坏采集卡"),
        ("安全 I/O", "急停、安全门、正反限位、驱动器就绪/报警的电气逻辑和端子图", "软件不能判定安全状态"),
    ], [1750, 4650, 2960], 8.1)
    add_callout(doc, "电流传感器特别提醒", "如果安全带电机是直流电机或无刷直流电机，普通交流电流互感器不能测量直流母线电流。应先明确测量的是直流母线、相电流还是输入交流，再选霍尔电流传感器、隔离电流变送器或分流器方案。", "danger")

    add_heading(doc, "4 硬件选型与电气设计建议", 1)
    add_heading(doc, "4.1 工控机", 2)
    add_bullets(doc, [
        "Windows 10/11 IoT Enterprise LTSC 或企业认可的长期支持版本；x64 驱动必须覆盖 CAN 卡和采集卡。",
        "至少预留 2 个可用 PCIe 插槽；购买前核对全高/半高、x1/x4 插槽和机箱空间。",
        "建议 16 GB 内存、工业 SSD 512 GB 以上，并预留独立数据分区或定期外部备份。",
        "双网口更方便隔离设备网与办公网；配置 UPS、可靠接地和散热。",
        "禁用自动睡眠、休眠和试验期间自动重启；Windows 更新采用维护窗口。"
    ])
    add_heading(doc, "4.2 CAN 卡与总线", 2)
    add_bullets(doc, [
        "优先选择 PCIe 工业 CAN 卡，带电气隔离、x64 Windows 驱动和明确的 C# 调用示例。",
        "总线两端各配置 120 Ω 终端电阻；断电测量 CAN_H 与 CAN_L 之间应约为 60 Ω。",
        "使用双绞屏蔽 CAN 电缆，主干拓扑，支线尽可能短；屏蔽层按电气设计单点或规范接地。",
        "确认标准帧/扩展帧、CAN 2.0/CAN FD、波特率、节点 ID 和心跳机制。"
    ])
    add_heading(doc, "4.3 模拟量采集与传感器", 2)
    add_table(doc, ["项目", "建议"], [
        ("通道", "至少 4 AI，拉力/电流/电压各占 1 路并保留 1 路；不要在软件中复用同一通道"),
        ("输入", "优先统一为隔离 4~20 mA；若采用 0~10 V，注意共模、压降、屏蔽和距离"),
        ("分辨率", "16 位起步；精度看整条测量链而不是只看 ADC 位数"),
        ("采样", "普通耐久趋势 50~100 ms 可用；若需捕捉冲击峰值，按机械动态重新计算采样率"),
        ("隔离", "工业现场优先通道间或组间隔离，至少输入与工控机侧隔离"),
        ("拉力", "mV/V 称重传感器必须配专用变送器/放大器，输出与采集卡匹配"),
    ], [1900, 7460], 9)
    add_heading(doc, "4.4 安全回路", 2)
    add_callout(doc, "不可替代", "急停、安全门、机械限位和驱动器 STO/使能切断必须构成独立硬件安全链。上位机只做监视、记录和附加停机请求，不能作为唯一安全措施。", "danger")

    add_heading(doc, "5 软件安装、配置与运行模式", 1)
    add_heading(doc, "5.1 开发环境", 2)
    add_steps(doc, [
        "在开发机安装 .NET 8 SDK 和 Visual Studio 2022（.NET 桌面开发工作负载）。",
        "打开 DurabilityTestingSystem.sln，执行 dotnet restore。",
        "执行 dotnet build DurabilityTestingSystem.csproj，要求 0 警告、0 错误。",
        "运行 dotnet run；首次启动会在 LocalAppData 创建数据库。"
    ])
    add_code(doc, "dotnet restore\ndotnet build DurabilityTestingSystem.csproj -c Release\ndotnet run")
    add_heading(doc, "5.2 system-profile.json", 2)
    add_body(doc, "发布目录中的配置决定运行模式和现场适配器。正式配置示例：")
    add_code(doc, '{\n  "Mode": "Production",\n  "ProfileName": "现场正式配置",\n  "AutoConnectOnStartup": true,\n  "HardwareAdapterAssembly": "DurabilityTestingSystem.HardwareAdapter.Site.dll",\n  "HardwareAdapterType": "DurabilityTestingSystem.HardwareAdapter.Site.SiteHardwarePlatform",\n  "Notes": "经联调和验收后启用"\n}')
    add_bullets(doc, [
        "开发验证可使用 --mode=Production 临时覆盖模式，不会修改 JSON。",
        "未填写适配器程序集和类型时，正式模式安全失败并禁止启动。",
        "不要把厂家 DLL 放进源码根目录；放入适配器项目的受控 lib 目录或按厂家许可部署。",
        "每次修改正式配置都要记录日期、修改人、原因和回退版本。"
    ])
    add_picture(doc, ROOT / "artifacts" / "captures" / "production-diagnostics.png",
                "图 2  Production 模式未配置硬件时的安全拦截（实际截图）",
                "正式模式硬件未配置时设备诊断页面")

    add_heading(doc, "6 现场硬件适配器开发", 1)
    add_body(doc, "现场代码不要直接改 UI。复制 HardwareAdapter.Template 为 DurabilityTestingSystem.HardwareAdapter.Site，在项目中引用 CAN 卡和采集卡厂家 DLL，然后实现 IHardwarePlatform。")
    add_heading(doc, "6.1 实现顺序", 2)
    add_steps(doc, [
        "只实现 ConnectAndSelfCheckAsync：打开设备、读取版本和状态，不允许电机运动。",
        "实现 ReadSampleAsync：先用标准源/砝码验证工程量，不接电机控制。",
        "实现 StopAsync：优先保证任何时刻都能安全请求停止和去使能。",
        "在无负载、低速条件下实现 BeginPullAsync 和 BeginReturnAsync。",
        "实现 Hold、Pause、Reset，并加入命令应答、超时、重连和报警映射。",
        "完成故障注入后，才允许 Health.CanStartTest=true。"
    ])
    add_heading(doc, "6.2 适配器健康状态规则", 2)
    add_table(doc, ["设备键", "Online 的必要条件", "失败处理"], [
        ("can", "卡已打开，通道已启动，驱动器心跳和状态有效", "停止发送运动命令，状态置 Fault/Disconnected"),
        ("analog", "设备在线，三通道有效，无断线/超量程，标定可用", "报警并阻止启动；运行中请求停机"),
        ("motor", "驱动器就绪、无报警、模式正确、未意外运动", "去使能/停机并记录报警码"),
        ("safety", "急停释放、安全门闭合、限位状态合理", "硬件链先动作，软件记录并禁止启动"),
    ], [1500, 4650, 3210], 8.5)
    add_callout(doc, "实现要求", "厂家 SDK 回调通常来自非 UI 线程。适配器内部必须处理线程安全、句柄释放和取消；不要直接从 SDK 回调更新 WinForms 控件。", "warn")

    add_heading(doc, "7 CAN 通讯调试", 1)
    add_heading(doc, "7.1 上电前检查", 2)
    add_bullets(doc, [
        "确认 CAN_H、CAN_L、参考地接线和屏蔽；电源关闭时测量总线电阻约 60 Ω。",
        "核对驱动器节点 ID 与波特率，确保总线上没有重复节点。",
        "先用厂家工具或 CAN 分析仪确认能收到心跳/状态，再运行上位机。",
        "明确驱动器是否需要 NMT、使能序列、控制字或特定握手。"
    ])
    add_heading(doc, "7.2 协议映射表必须包含", 2)
    add_table(doc, ["字段", "示例/说明", "必须验证"], [
        ("帧类型", "标准帧或扩展帧", "SDK 标志与驱动器一致"),
        ("报文 ID", "命令、状态、心跳、报警分别记录", "方向和节点偏移正确"),
        ("字节序", "Little/Big Endian", "用已知值对照抓包"),
        ("比例与偏置", "raw × scale + offset", "正负号、单位和溢出"),
        ("超时", "如 1000 ms 无心跳", "运行中必须进入故障并停机"),
        ("确认机制", "回读状态或应答帧", "不能只发送不验证"),
    ], [1700, 3900, 3760], 8.6)
    add_heading(doc, "7.3 点动调试", 2)
    add_steps(doc, [
        "拆除或释放负载，设置最低安全速度和最小行程。",
        "只发送使能，确认电机不自行运动；读取实际状态。",
        "执行极短正向点动，核对机械方向、编码器方向和正限位。",
        "执行停止并测量停止时间；断开 CAN 验证驱动器自身超时策略。",
        "执行反向点动，核对原点/反限位。",
        "逐步增加速度和行程，每一级都保存抓包、日志和结果。"
    ])

    add_heading(doc, "8 模拟量采集、换算与标定", 1)
    add_heading(doc, "8.1 信号换算", 2)
    add_body(doc, "适配器应先把厂家原始码转换为电压/电流，再转换为工程量。典型公式如下：")
    add_code(doc, "4~20 mA：工程量 = (I_mA - 4) / 16 × 满量程\n0~10 V： 工程量 = V / 10 × 满量程\n两点标定：工程量 = a × 原始量 + b")
    add_bullets(doc, [
        "4~20 mA 低于约 3.6 mA 或超出合理范围时，应判为断线/故障，而不是继续显示负值。",
        "0~10 V 输入必须确认采集卡共地/差分方式和允许共模电压。",
        "拉力有方向时要确认正负号；当前 UI 默认以拉力正值显示。",
        "滤波只能降低噪声，不能掩盖真实冲击峰值；保护通道可使用比显示通道更快的路径。"
    ])
    add_heading(doc, "8.2 拉力标定步骤", 2)
    add_steps(doc, [
        "机械空载并稳定预热，执行零点采集，记录原始值 X0。",
        "施加经过溯源的标准载荷 F1，稳定后记录 X1；再取 50%~80% 量程点 F2/X2。",
        "计算线性系数 a 与 b，检查回零、重复性和加载/卸载回差。",
        "用至少 3 个非标定点复核误差；误差超限时检查机械安装、变送器和接线。",
        "保存标定日期、设备序列号、标准器编号、人员、温度和系数；锁定生产参数。"
    ])
    add_heading(doc, "8.3 电流与电压验证", 2)
    add_bullets(doc, [
        "使用钳形表/标准表与上位机同时读数，覆盖零点、常用点和接近保护值的点。",
        "确认电流方向、直流/交流类型、RMS 或平均值定义及采样带宽。",
        "不要直接把电机高电压或大电流接入采集卡；必须经隔离传感器/变送器。",
        "把硬件极限设定在设备安全范围内；软件报警值应留出传感器误差和动态裕量。"
    ])
    add_picture(doc, ROOT / "artifacts" / "captures" / "settings.png",
                "图 3  当前参数设置页：CAN、模拟量通道、量程与安全输入（实际截图）",
                "参数设置页面截图")

    add_heading(doc, "9 安全联锁与保护", 1)
    add_heading(doc, "9.1 分层保护", 2)
    add_table(doc, ["层级", "内容", "责任"], [
        ("机械", "防护罩、机械限位、强度裕量、防飞出", "机械设计"),
        ("电气", "急停、安全继电器、STO/使能切断、空开/熔断", "电气设计"),
        ("驱动器", "过流、过压、过温、堵转、通讯超时", "驱动器参数"),
        ("上位机", "阈值、状态一致性、通讯监视、报警记录", "软件"),
    ], [1300, 5050, 3010], 9)
    add_heading(doc, "9.2 必做故障注入", 2)
    add_table(doc, ["故障", "预期硬件动作", "预期软件表现"], [
        ("按下急停", "立即切断危险运动", "安全联锁 Fault，运行中报警，禁止再次启动"),
        ("打开安全门", "按风险评估停止/禁止运动", "记录时间与状态，CanStartTest=false"),
        ("触发正/反限位", "相应方向不能继续运动", "报警并允许受控反向脱离（需工艺确认）"),
        ("拔掉 CAN", "驱动器按超时策略停止", "CAN Fault，停止采样/运动流程并保存异常记录"),
        ("断开拉力传感器", "硬件保护保持有效", "模拟量 Fault，显示断线而不是 0 N"),
        ("制造超拉力", "硬件极限先保护", "软件阈值报警并请求停止，保存峰值"),
    ], [1900, 3550, 3910], 8.3)
    add_callout(doc, "验收红线", "任何一次急停、限位、断线或通讯故障如果仍能继续发送运动命令，均不得进入带载试验和交付。", "danger")

    add_heading(doc, "10 配方、试验与数据操作", 1)
    add_heading(doc, "10.1 新建/修改方案", 2)
    add_steps(doc, [
        "进入“试验方案”，选择已有方案或新建方案。",
        "填写方案编号、名称、循环次数和目标拉力。",
        "按工艺添加正向拉伸、负载保持、反向回程、等待和循环计数步骤。",
        "保存后重新选择该方案，确认步骤、时长和完成条件已正确回读。",
        "正式试验前由工艺/质量人员批准方案版本。"
    ])
    add_picture(doc, ROOT / "artifacts" / "captures" / "plans-v1.png",
                "图 4  方案和循环步骤编辑页（实际截图）", "试验方案页面截图")
    add_heading(doc, "10.2 日常试验 SOP", 2)
    add_steps(doc, [
        "检查防护、急停、限位、夹具和试件安装；清空危险区域。",
        "启动软件并进入“设备诊断”，执行连接与自检。",
        "进入“参数设置”，只核对已批准的量程、通道、速度和保护值；不要现场随意改保护。",
        "进入“试验控制”，选择方案并填写唯一试件编号。",
        "低速预运行 1~3 个循环，确认方向、拉力、电流、电压和回零。",
        "启动正式循环，观察前若干循环；无人值守条件必须由项目风险评估决定。",
        "正常结束后核对结果和峰值，执行数据库备份或班次自动备份。"
    ])
    add_heading(doc, "10.3 历史数据", 2)
    add_body(doc, "历史页支持日期、关键词和结果筛选，导出 CSV。正式报告模板尚需甲方确认；SQLite 中已分别保存试验汇总与采样点，避免只保存最终峰值。")
    add_picture(doc, ROOT / "artifacts" / "captures" / "history.png",
                "图 5  历史记录筛选与 CSV 导出（实际截图）", "历史数据页面截图")

    add_heading(doc, "11 分阶段现场调试流程", 1)
    add_heading(doc, "阶段 A：离线软件验证", 2)
    add_bullets(doc, [
        "Demo 模式检查所有页面、配方保存、筛选、CSV、日志和数据库备份。",
        "执行至少 30 分钟模拟运行，确认 UI 无卡死、内存无持续异常增长。",
        "记录显示分辨率和 DPI；在现场工控机上复测页面无裁切。"
    ])
    add_heading(doc, "阶段 B：设备单体连接（禁止运动）", 2)
    add_bullets(doc, [
        "安装厂家驱动，记录版本；设备管理器无异常。",
        "CAN 卡可打开/关闭；采集卡可读取；安全输入能逐点变化。",
        "Production 诊断页能显示每个设备的真实状态和明确错误。"
    ])
    add_heading(doc, "阶段 C：传感器标定", 2)
    add_bullets(doc, [
        "拉力、电流、电压分别完成零点、量程和多点复核。",
        "保存原始值、工程量、误差、标定工具和证书信息。",
        "断线、短路和超量程均能被识别。"
    ])
    add_heading(doc, "阶段 D：无负载低速运动", 2)
    add_bullets(doc, [
        "确认方向、位置、正反限位、停止、急停和通讯丢失。",
        "确认 Pause 与 Stop 的机械行为符合风险评估。",
        "抓取 CAN 报文并与协议表逐项对照。"
    ])
    add_heading(doc, "阶段 E：小负载联动", 2)
    add_bullets(doc, [
        "从额定目标的 10%~20% 开始，逐级提高负载。",
        "每一级检查拉力、电流、温升、机械噪声和停止距离。",
        "验证配方阶段切换与采样时间戳。"
    ])
    add_heading(doc, "阶段 F：额定与耐久验证", 2)
    add_bullets(doc, [
        "执行额定工况、连续运行、断电恢复、数据库增长和备份恢复测试。",
        "完成甲方见证的故障注入和验收矩阵。",
        "冻结软件、适配器、驱动、配置和电气图版本。"
    ])

    add_heading(doc, "12 验收测试", 1)
    add_table(doc, ["编号", "测试项", "方法", "通过标准"], [
        ("SW-01", "模式隔离", "分别启动 Demo/Production", "数据库分离；正式未自检时禁止启动"),
        ("SW-02", "参数校验", "设置重复通道、错误量程/阈值", "保存或启动被阻止并给出原因"),
        ("HW-01", "CAN 稳定性", "连续运行并统计错误帧/超时", "达到项目约定时长且无失控"),
        ("HW-02", "采集准确度", "标准源/标准载荷多点比对", "满足项目精度指标"),
        ("SAFE-01", "急停", "运行中按急停", "硬件立即停；软件报警且不可重启"),
        ("SAFE-02", "限位", "低速触发正反限位", "危险方向禁止，状态和日志正确"),
        ("DATA-01", "数据完整性", "完成/停止/报警三类试验", "汇总、采样、时间、配方和日志一致"),
        ("DATA-02", "备份恢复", "备份后在测试机恢复", "数据库 integrity_check=ok，记录可查询"),
        ("PERF-01", "长稳", "额定采样连续运行", "无卡死、无持续内存增长、无明显丢样"),
    ], [1000, 1700, 3300, 3360], 8.0)
    add_body(doc, "建议把上表复制为正式 FAT/SAT 记录，增加实际结果、证据文件、测试人、见证人和签字日期。")

    add_heading(doc, "13 数据库、备份与恢复", 1)
    add_heading(doc, "13.1 数据表", 2)
    add_table(doc, ["表", "内容", "注意"], [
        ("settings", "当前试验与硬件参数 JSON", "修改会写系统日志"),
        ("plans / plan_steps", "方案主表与动作步骤", "方案版本化可在二期扩展"),
        ("test_records", "每次试验汇总", "试验号应唯一"),
        ("test_samples", "采样点", "数据量最大，需制定保留/归档策略"),
        ("system_logs", "系统、操作、通讯和报警日志", "不应记录密码或敏感凭据"),
    ], [1900, 4200, 3260], 8.6)
    add_heading(doc, "13.2 备份", 2)
    add_steps(doc, [
        "在“设备诊断”点击“备份数据库”，软件使用 SQLite 在线备份 API。",
        "把 Backups 文件夹同步到另一块物理磁盘或受控服务器；同盘备份不能防磁盘损坏。",
        "每班或每天备份，按质量要求制定保留周期。",
        "每月至少执行一次恢复演练，而不是只检查备份文件存在。"
    ])
    add_heading(doc, "13.3 恢复", 2)
    add_steps(doc, [
        "停止试验并退出软件，备份当前数据库文件。",
        "把待恢复备份复制为 durability.db，保留原文件时间戳和恢复记录。",
        "启动软件进入设备诊断，确认数据库完整性为 ok。",
        "抽查参数、方案、历史记录和采样点；完成后记录恢复人和原因。"
    ])

    add_heading(doc, "14 故障排查", 1)
    add_table(doc, ["现象", "优先检查", "处理"], [
        ("CAN 未配置", "system-profile.json 的 DLL 和类型", "构建适配器并复制依赖；检查完整类型名"),
        ("CAN 离线", "驱动、通道、波特率、节点、60 Ω", "先用厂家工具验证，再查 SDK 错误码"),
        ("能发不能收", "ID、过滤器、标准/扩展帧、接线", "关闭过滤或抓包对比协议"),
        ("拉力恒为 0", "传感器供电、变送输出、通道、断线", "用万用表测信号，不要先改比例"),
        ("拉力漂移", "预热、接地、机械应力、滤波", "重新零点，排除安装与温漂"),
        ("电流明显错误", "是否用交流互感器测直流", "按被测波形改用合适的霍尔/变送器"),
        ("启动被阻止", "设备诊断摘要和日志", "修复所有关键设备，不要绕过 CanStartTest"),
        ("数据库异常", "磁盘空间、权限、完整性", "停止试验，备份现状，从有效备份恢复"),
        ("页面裁切", "Windows 缩放、分辨率、DPI", "使用 1920×1080/100% 或现场验证过的设置"),
    ], [2200, 3550, 3610], 8.0)
    add_picture(doc, ROOT / "artifacts" / "captures" / "diagnostics-v1.png",
                "图 6  设备诊断、数据库完整性和备份入口（实际截图）", "设备诊断页面截图")

    add_heading(doc, "15 运维、变更与正式交付", 1)
    add_heading(doc, "15.1 发布步骤", 2)
    add_code(doc, "dotnet publish DurabilityTestingSystem.csproj -c Release -r win-x64 --self-contained true -o publish\\win-x64")
    add_steps(doc, [
        "在干净测试机安装厂家驱动，复制发布目录和现场适配器 DLL/依赖。",
        "配置 system-profile.json 为 Production；不要覆盖现场数据库。",
        "运行设备诊断和 FAT/SAT；通过后计算发布包哈希并归档。",
        "建立桌面快捷方式，设置应用数据目录备份和维护权限。",
        "冻结软件版本、适配器版本、厂家驱动版本、配置文件、电气图和协议文档。"
    ])
    add_heading(doc, "15.2 变更控制", 2)
    add_bullets(doc, [
        "任何 CAN 报文、量程、标定、保护阈值和安全逻辑变更都必须重新回归。",
        "不要在试验运行中替换 DLL、配置或数据库。",
        "保留上一个可运行发布包和数据库备份，确保可回退。",
        "二次开发优先扩展现场适配器，不要把厂家代码散布到页面层。"
    ])
    add_heading(doc, "15.3 最终交付清单", 2)
    add_bullets(doc, [
        "源代码、Release 发布包、现场硬件适配器源码与 DLL；",
        "system-profile.json、参数清单、配方清单和标定系数；",
        "CAN 协议、厂家 SDK/驱动及版本说明；",
        "电气原理图、端子图、I/O 表、传感器说明书和标定证书；",
        "FAT/SAT、故障注入、长稳和备份恢复记录；",
        "本手册、操作员培训记录、维护联系人和版本变更记录。"
    ])
    add_callout(doc, "正式启用判据", "只有当硬件资料冻结、适配器完成、设备诊断全部在线、安全故障注入通过、数据准确度通过并完成甲方验收后，才可把该系统作为正式试验设备使用。", "ok")

    doc.add_page_break()
    add_heading(doc, "附录 A 现场开机检查表", 1)
    add_table(doc, ["检查项", "结果/记录"], [
        ("防护罩、夹具、试件安装可靠", "□ 正常  □ 异常：____________"),
        ("急停释放且硬件安全回路正常", "□ 正常  □ 异常：____________"),
        ("正反限位和安全门输入正确", "□ 正常  □ 异常：____________"),
        ("CAN、模拟量、电机、安全联锁全部在线", "□ 正常  □ 异常：____________"),
        ("方案、试件号、量程和保护参数已核对", "□ 正常  □ 异常：____________"),
        ("低速预运行通过", "□ 正常  □ 异常：____________"),
        ("数据库备份策略可用", "□ 正常  □ 异常：____________"),
        ("操作员：________  日期：________  复核：________", ""),
    ], [5450, 3910], 9)

    add_heading(doc, "附录 B 当前尚待甲方/采购确认的问题", 1)
    add_compact_bullets(doc, [
        "安全带电机及驱动器准确型号、供电、控制模式和 CAN 协议；",
        "CAN 卡品牌型号、通道数、隔离、PCIe 规格、SDK 与驱动版本；",
        "模拟量采集卡接口形式、输入类型、通道数、采样率、同步性和 SDK；",
        "拉力传感器量程、精度、输出、变送器和标定要求；",
        "电流测量点与波形类型，确认普通电流互感器是否适用；",
        "电压测量点、最高电压、隔离和输出比例；",
        "急停、安全门、限位、STO/使能回路及安全等级要求；",
        "目标采样率、数据保留周期、报告格式、用户权限和审计要求；",
        "正式验收的精度、节拍、连续运行时长与判定标准。"
    ])

    doc.core_properties.title = "安全带耐久试验系统 软件使用调试与落地指导手册"
    doc.core_properties.subject = "工控机上位机、CAN、模拟量采集、现场调试与验收"
    doc.core_properties.author = "沈阳艾德瑞自动化有限公司"
    doc.core_properties.comments = "基于 DurabilityTestingSystem V1.0 软件基线生成"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
